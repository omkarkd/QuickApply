"""
Document text extraction for the Upload page.

Supported formats (case-insensitive): .pdf, .docx, .doc, .txt

Each public function takes either a file path (str | Path) or a binary
blob (bytes) plus a filename, and returns the extracted plain text as a str.

The module is intentionally Streamlit-free so it can be unit-tested with
plain pytest / a REPL.
"""
from __future__ import annotations

import io
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Union

PathOrBytes = Union[str, bytes, Path]


class UnsupportedFileType(ValueError):
    pass


# ---------- Public API ----------

def extract(path: PathOrBytes, *, filename: str | None = None) -> str:
    """Dispatch to the right extractor based on file extension."""
    name = filename or (str(path) if isinstance(path, (str, Path)) else "")
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".doc":
        return extract_doc(path)
    if suffix == ".txt":
        return extract_txt(path)
    raise UnsupportedFileType(
        f"Unsupported file type '{suffix}'. Allowed: .pdf, .docx, .doc, .txt"
    )


def extract_pdf(src: PathOrBytes) -> str:
    """Extract text from a PDF using pypdf."""
    from pypdf import PdfReader  # local import — only needed for this branch

    reader = PdfReader(_to_fileobj(src))
    chunks: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001
            # Don't kill the whole upload for one bad page; log and continue.
            text = f"[page {i + 1} extraction error: {e}]"
        if text:
            chunks.append(text)
    return "\n\n".join(chunks).strip()


def extract_docx(src: PathOrBytes) -> str:
    """Extract text from a DOCX using python-docx (paragraphs + table cells)."""
    from docx import Document  # python-docx

    doc = Document(_to_fileobj(src))
    pieces: list[str] = []

    for para in doc.paragraphs:
        if para.text:
            pieces.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells if cell.text)
            if row_text:
                pieces.append(row_text)

    return "\n".join(pieces).strip()


def extract_doc(src: PathOrBytes) -> str:
    """
    Extract text from a legacy .doc (OLE compound doc) by shelling out to the
    `antiword` binary. We:
      1. Write the bytes to a temp file (antiword needs a path)
      2. Run `antiword <tmp>`
      3. Return stdout
    """
    if shutil.which("antiword") is None:
        raise RuntimeError(
            "antiword binary not found on PATH. Install with: brew install antiword"
        )

    blob = _to_bytes(src)
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as tmp:
        tmp.write(blob)
        tmp.flush()
        try:
            proc = subprocess.run(
                ["antiword", tmp.name],
                capture_output=True,
                text=True,
                timeout=60,
            )
        except subprocess.TimeoutExpired as e:
            raise RuntimeError("antiword timed out reading the .doc") from e

    if proc.returncode != 0:
        raise RuntimeError(
            f"antiword failed (exit {proc.returncode}): {proc.stderr.strip()[:200]}"
        )
    return proc.stdout.strip()


def extract_txt(src: PathOrBytes) -> str:
    """Read a text file. Tries utf-8 first, then latin-1 as a permissive fallback."""
    raw = _to_bytes(src)
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    # Last resort — replace undecodable bytes
    return raw.decode("utf-8", errors="replace").strip()


# ---------- Helpers ----------

def _to_bytes(src: PathOrBytes) -> bytes:
    if isinstance(src, bytes):
        return src
    return Path(src).read_bytes()


def _to_fileobj(src: PathOrBytes):
    """Return either a BytesIO or a path string — pypdf/python-docx accept both."""
    if isinstance(src, (str, Path)):
        return str(src)
    return io.BytesIO(src)
