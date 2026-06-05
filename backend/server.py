import os
import logging
import re
import io
import json
import asyncio
import anthropic
from pathlib import Path
from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, Request, Response
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo.errors import DuplicateKeyError
from pydantic import BaseModel, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib import colors
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ConversationHandler, ContextTypes

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Telegram Bot States
CHOOSING, WAITING_FOR_TEXT = range(2)

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Session duration
SESSION_DURATION_DAYS = 7

# Create the main app
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ==================== Models ====================

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    created_at: str

class ResumeData(BaseModel):
    name: Optional[str] = ""
    email: Optional[str] = ""
    linkedin: Optional[str] = ""
    professional_summary: Optional[str] = ""
    technical_skills: List[str] = []
    experiences: List[dict] = []
    projects: List[dict] = []
    education: List[dict] = []

class ResumeCreate(BaseModel):
    title: str
    raw_text: str
    parsed_data: Optional[ResumeData] = None

class ResumeUpdate(BaseModel):
    title: Optional[str] = None
    raw_text: Optional[str] = None
    parsed_data: Optional[ResumeData] = None

class ResumeResponse(BaseModel):
    id: str
    user_id: str
    title: str
    raw_text: str
    parsed_data: Optional[dict] = None
    created_at: str
    updated_at: str

class ParseRequest(BaseModel):
    text: str

# JD (Job Description) Models
class JDData(BaseModel):
    job_title: Optional[str] = ""
    company: Optional[str] = ""
    location: Optional[str] = ""
    must_have_skills: List[str] = []
    good_to_have_skills: List[str] = []
    experience_required: Optional[str] = ""
    salary_range: Optional[str] = ""
    job_type: Optional[str] = ""  # Full-time/Part-time/Contract
    description: Optional[str] = ""

class JDCreate(BaseModel):
    title: str
    raw_text: str
    parsed_data: Optional[JDData] = None

class JDUpdate(BaseModel):
    title: Optional[str] = None
    raw_text: Optional[str] = None
    parsed_data: Optional[JDData] = None

class JDResponse(BaseModel):
    id: str
    user_id: str
    title: str
    raw_text: str
    parsed_data: Optional[dict] = None
    created_at: str
    updated_at: str

# Telegram Settings Model
class TelegramSettingsUpdate(BaseModel):
    bot_token: Optional[str] = None
    is_active: Optional[bool] = None

# Uploaded Texts (intermediate store for the Upload Document page)
class UploadedTextCreate(BaseModel):
    doc_type: str  # "resume" | "jd"
    filename: str
    raw_text: str

class UploadedTextResponse(BaseModel):
    id: str
    user_id: str
    doc_type: str
    filename: str
    raw_text: str
    char_count: int
    created_at: str

# Resume Rewriter Models
class RewriteRequest(BaseModel):
    resume_text: str
    jd_text: str

# Match Score Models
class MatchScoreRequest(BaseModel):
    resume_text: str
    jd_text: str

class KeywordDetail(BaseModel):
    keyword: str
    present: bool
    context: Optional[str] = None

class MatchScoreResponse(BaseModel):
    overall_score: float
    keywords_score: float
    alignment_score: float
    rpm_score: float
    present_keywords: List[str]
    missing_keywords: List[str]
    keyword_details: List[KeywordDetail]
    suggestions: List[str]

# ==================== Auth Helpers ====================

# Keep password auth for backward compatibility
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

async def get_current_user(request: Request):
    """Get current user from session token (cookie or header)"""
    # Try cookie first
    session_token = request.cookies.get("session_token")
    
    # Fallback to Authorization header
    if not session_token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            session_token = auth_header.split(" ")[1]
    
    if not session_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    # Check session in database
    session = await db.user_sessions.find_one({"session_token": session_token}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=401, detail="Invalid session")
    
    # Check expiry with timezone awareness
    expires_at = session.get("expires_at")
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Session expired")
    
    # Get user
    user = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    
    return user

# ==================== Auth Routes ====================

@api_router.post("/auth/register")
async def register(user_data: UserCreate, response: Response):
    """Email/password registration. Backed by MongoDB with a unique-email index.

    The unique index (created in init_auth_indexes) is the source of truth for
    duplicate prevention; the find_one precheck is a fast-path for the common
    case. If a race lets two requests through, the second insert raises
    DuplicateKeyError and we translate it to a 400.
    """
    if len(user_data.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    if len(user_data.password) > 128:
        raise HTTPException(status_code=400, detail="Password too long")
    if not user_data.name.strip():
        raise HTTPException(status_code=400, detail="Name is required")

    # Fast-path duplicate check
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 1})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")

    user_id = f"user_{uuid.uuid4().hex[:12]}"
    user_doc = {
        "user_id": user_id,
        "email": user_data.email,
        "name": user_data.name.strip(),
        "password": hash_password(user_data.password),
        "picture": None,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    try:
        await db.users.insert_one(user_doc)
    except DuplicateKeyError:
        # The unique index caught a race between the find_one and the insert
        raise HTTPException(status_code=400, detail="Email already registered")
    except Exception as e:  # noqa: BLE001
        logger.error(f"register insert failed: {e}")
        raise HTTPException(status_code=500, detail="Could not create account")

    # Create session
    session_token = f"session_{uuid.uuid4().hex}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=SESSION_DURATION_DAYS)

    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at,
        "created_at": datetime.now(timezone.utc)
    })

    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=SESSION_DURATION_DAYS * 24 * 60 * 60
    )

    return {
        "user": UserResponse(
            user_id=user_id,
            email=user_data.email,
            name=user_doc["name"],
            picture=None,
            created_at=user_doc["created_at"]
        ),
        "session_token": session_token
    }

@api_router.post("/auth/login")
async def login(login_data: UserLogin, response: Response):
    """Legacy email/password login"""
    user = await db.users.find_one({"email": login_data.email}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check if user has password
    if not user.get("password"):
        raise HTTPException(status_code=401, detail="This account has no password set")

    if not verify_password(login_data.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Create session
    session_token = f"session_{uuid.uuid4().hex}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=SESSION_DURATION_DAYS)
    
    await db.user_sessions.insert_one({
        "user_id": user["user_id"],
        "session_token": session_token,
        "expires_at": expires_at,
        "created_at": datetime.now(timezone.utc)
    })
    
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=SESSION_DURATION_DAYS * 24 * 60 * 60
    )
    
    return {
        "user": UserResponse(
            user_id=user["user_id"],
            email=user["email"],
            name=user["name"],
            picture=user.get("picture"),
            created_at=user["created_at"]
        ),
        "session_token": session_token
    }

@api_router.get("/auth/me")
async def get_me(request: Request):
    """Get current user from session"""
    user = await get_current_user(request)
    return UserResponse(
        user_id=user["user_id"],
        email=user["email"],
        name=user["name"],
        picture=user.get("picture"),
        created_at=user["created_at"]
    )

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    """Logout and clear session.

    The session token can arrive via:
      1. The httpOnly cookie (browser path), or
      2. The Authorization: Bearer header (API clients like the Streamlit UI)

    Both paths must be honored — otherwise a Bearer-authenticated logout would
    silently return 200 without invalidating the server-side session.
    """
    session_token = request.cookies.get("session_token")
    if not session_token:
        auth_header = request.headers.get("Authorization", "")
        if auth_header.startswith("Bearer "):
            session_token = auth_header.split(" ", 1)[1].strip()

    if session_token:
        await db.user_sessions.delete_one({"session_token": session_token})

    response.delete_cookie(key="session_token", path="/")
    return {"message": "Logged out successfully"}

# Helper to get user from request for route dependencies
async def get_user_dependency(request: Request):
    return await get_current_user(request)

# ==================== LLM Helper ====================

# Configure Anthropic SDK to talk to the local model gateway.
# By default: opencode.ai/zen with model `minimax-m3-free`.
# Override via env vars ANTHROPIC_BASE_URL, ANTHROPIC_MODEL, ANTHROPIC_API_KEY.
ANTHROPIC_BASE_URL = os.environ.get("ANTHROPIC_BASE_URL", "https://opencode.ai/zen")
ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_MODEL", "minimax-m3-free")

anthropic_client = None


def get_anthropic_client() -> anthropic.Anthropic:
    """Lazy-init the Anthropic client so missing key fails fast with a clear error."""
    global anthropic_client
    if anthropic_client is None:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=500,
                detail="ANTHROPIC_API_KEY not configured. Set it in backend/.env",
            )
        anthropic_client = anthropic.Anthropic(
            api_key=api_key,
            base_url=ANTHROPIC_BASE_URL,
        )
    return anthropic_client


async def llm_complete_json(system_prompt: str, user_prompt: str) -> dict:
    """Call the LLM and return the parsed JSON response.

    Strips markdown code fences if the LLM wraps the JSON in them.
    Raises HTTPException(500) on parse failure.
    """
    client = get_anthropic_client()
    try:
        message = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=4096,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
    except anthropic.APIError as e:
        logger.error(f"LLM API error: {e}")
        raise HTTPException(status_code=500, detail=f"LLM call failed: {str(e)}")

    # Concatenate all text blocks from the response
    raw = "".join(
        block.text for block in message.content if getattr(block, "type", None) == "text"
    )

    # Strip markdown fences
    json_str = raw
    if "```json" in json_str:
        json_str = json_str.split("```json", 1)[1].split("```", 1)[0]
    elif "```" in json_str:
        json_str = json_str.split("```", 1)[1].split("```", 1)[0]

    try:
        return json.loads(json_str.strip())
    except json.JSONDecodeError:
        logger.error(f"Failed to parse LLM JSON: {raw[:500]}")
        raise HTTPException(status_code=500, detail="Failed to parse AI response")


# ==================== Parse Routes ====================

@api_router.post("/parse/simple")
async def parse_simple(request: ParseRequest):
    """Simple regex-based resume parsing"""
    text = request.text
    
    # Extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    email = email_match.group(0) if email_match else ""
    
    # Extract LinkedIn
    linkedin_match = re.search(r'(?:linkedin\.com/in/|linkedin:?\s*)([^\s,\n]+)', text, re.IGNORECASE)
    linkedin = linkedin_match.group(0) if linkedin_match else ""
    
    # Extract name (first line or first capitalized words)
    lines = text.strip().split('\n')
    name = lines[0].strip() if lines else ""
    
    # Extract sections based on common headers
    sections = {
        "professional_summary": "",
        "technical_skills": [],
        "experiences": [],
        "projects": [],
        "education": []
    }
    
    current_section = None
    section_content = []
    
    section_keywords = {
        "summary": "professional_summary",
        "professional summary": "professional_summary",
        "about": "professional_summary",
        "profile": "professional_summary",
        "objective": "professional_summary",
        "skills": "technical_skills",
        "technical skills": "technical_skills",
        "technologies": "technical_skills",
        "experience": "experiences",
        "work experience": "experiences",
        "employment": "experiences",
        "project": "projects",
        "projects": "projects",
        "education": "education",
        "academic": "education",
        "qualification": "education"
    }
    
    for line in lines[1:]:
        line_lower = line.lower().strip()
        line_clean = re.sub(r'[:\-]', '', line_lower).strip()
        
        # Check if this is a section header
        matched_section = None
        for keyword, section in section_keywords.items():
            if keyword in line_clean and len(line_clean) < 50:
                matched_section = section
                break
        
        if matched_section:
            # Save previous section
            if current_section and section_content:
                if current_section == "professional_summary":
                    sections[current_section] = '\n'.join(section_content)
                elif current_section == "technical_skills":
                    # Parse skills as comma/bullet separated
                    all_skills = '\n'.join(section_content)
                    skills = re.split(r'[,\n•\-\|]', all_skills)
                    sections[current_section] = [s.strip() for s in skills if s.strip()]
                elif current_section in ["experiences", "projects", "education"]:
                    # Group by bullet points
                    items = []
                    current_item = {"title": "", "description": "", "bullets": []}
                    for content_line in section_content:
                        if content_line.startswith(('•', '-', '*')) or re.match(r'^\d+\.', content_line):
                            bullet = re.sub(r'^[•\-\*\d+\.]\s*', '', content_line)
                            current_item["bullets"].append(bullet)
                        elif current_item["bullets"] or current_item["title"]:
                            if current_item["title"]:
                                items.append(current_item)
                            current_item = {"title": content_line, "description": "", "bullets": []}
                        else:
                            current_item["title"] = content_line
                    if current_item["title"] or current_item["bullets"]:
                        items.append(current_item)
                    sections[current_section] = items
            
            current_section = matched_section
            section_content = []
        elif current_section and line.strip():
            section_content.append(line.strip())
    
    # Save last section
    if current_section and section_content:
        if current_section == "professional_summary":
            sections[current_section] = '\n'.join(section_content)
        elif current_section == "technical_skills":
            all_skills = '\n'.join(section_content)
            skills = re.split(r'[,\n•\-\|]', all_skills)
            sections[current_section] = [s.strip() for s in skills if s.strip()]
        elif current_section in ["experiences", "projects", "education"]:
            items = []
            current_item = {"title": "", "description": "", "bullets": []}
            for content_line in section_content:
                if content_line.startswith(('•', '-', '*')) or re.match(r'^\d+\.', content_line):
                    bullet = re.sub(r'^[•\-\*\d+\.]\s*', '', content_line)
                    current_item["bullets"].append(bullet)
                elif current_item["bullets"] or current_item["title"]:
                    if current_item["title"]:
                        items.append(current_item)
                    current_item = {"title": content_line, "description": "", "bullets": []}
                else:
                    current_item["title"] = content_line
            if current_item["title"] or current_item["bullets"]:
                items.append(current_item)
            sections[current_section] = items
    
    return {
        "name": name,
        "email": email,
        "linkedin": linkedin,
        **sections
    }

@api_router.post("/parse/ai")
async def parse_ai(request: ParseRequest):
    """AI-powered resume parsing using Claude"""
    try:
        system_prompt = """You are a professional resume parser. Extract structured data from resume text.
            Return ONLY valid JSON with this exact structure:
            {
                "name": "Full Name",
                "email": "email@example.com",
                "linkedin": "linkedin.com/in/profile",
                "professional_summary": "Summary text",
                "technical_skills": ["skill1", "skill2"],
                "experiences": [
                    {
                        "title": "Job Title at Company",
                        "company": "Company Name",
                        "from_date": "Start Date",
                        "to_date": "End Date",
                        "description": "Brief description",
                        "bullets": ["Achievement 1", "Achievement 2"]
                    }
                ],
                "projects": [
                    {
                        "title": "Project Name",
                        "description": "Project description",
                        "bullets": ["Feature 1", "Feature 2"]
                    }
                ],
                "education": [
                    {
                        "title": "Degree - Institution",
                        "institution": "Institution Name",
                        "from_date": "Start Date",
                        "to_date": "End Date",
                        "description": "Additional info"
                    }
                ]
            }
            Extract all information accurately. If something is not found, use empty string or empty array."""
        return await llm_complete_json(
            system_prompt=system_prompt,
            user_prompt=f"Parse this resume and return JSON:\n\n{request.text}",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI parsing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"AI parsing failed: {str(e)}")

# ==================== Resume CRUD Routes ====================

@api_router.post("/resumes", response_model=ResumeResponse)
async def create_resume(resume_data: ResumeCreate, current_user: dict = Depends(get_user_dependency)):
    resume_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    resume_doc = {
        "id": resume_id,
        "user_id": current_user["user_id"],
        "title": resume_data.title,
        "raw_text": resume_data.raw_text,
        "parsed_data": resume_data.parsed_data.model_dump() if resume_data.parsed_data else None,
        "created_at": now,
        "updated_at": now
    }
    
    await db.resumes.insert_one(resume_doc)
    
    return ResumeResponse(**{k: v for k, v in resume_doc.items() if k != "_id"})

@api_router.get("/resumes", response_model=List[ResumeResponse])
async def list_resumes(current_user: dict = Depends(get_user_dependency)):
    resumes = await db.resumes.find(
        {"user_id": current_user["user_id"]}, 
        {"_id": 0}
    ).sort("updated_at", -1).to_list(100)
    return [ResumeResponse(**r) for r in resumes]

@api_router.get("/resumes/{resume_id}", response_model=ResumeResponse)
async def get_resume(resume_id: str, current_user: dict = Depends(get_user_dependency)):
    resume = await db.resumes.find_one(
        {"id": resume_id, "user_id": current_user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return ResumeResponse(**resume)

@api_router.put("/resumes/{resume_id}", response_model=ResumeResponse)
async def update_resume(resume_id: str, resume_data: ResumeUpdate, current_user: dict = Depends(get_user_dependency)):
    resume = await db.resumes.find_one({"id": resume_id, "user_id": current_user["user_id"]})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    update_dict = {}
    if resume_data.title is not None:
        update_dict["title"] = resume_data.title
    if resume_data.raw_text is not None:
        update_dict["raw_text"] = resume_data.raw_text
    if resume_data.parsed_data is not None:
        update_dict["parsed_data"] = resume_data.parsed_data.model_dump()
    
    update_dict["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.resumes.update_one({"id": resume_id}, {"$set": update_dict})
    
    updated_resume = await db.resumes.find_one({"id": resume_id}, {"_id": 0})
    return ResumeResponse(**updated_resume)

@api_router.delete("/resumes/{resume_id}")
async def delete_resume(resume_id: str, current_user: dict = Depends(get_user_dependency)):
    result = await db.resumes.delete_one({"id": resume_id, "user_id": current_user["user_id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {"message": "Resume deleted successfully"}

# ==================== Uploaded Texts (intermediate store) ====================

VALID_DOC_TYPES = ("resume", "jd")

@api_router.post("/uploads", response_model=UploadedTextResponse)
async def create_upload(payload: UploadedTextCreate, current_user: dict = Depends(get_user_dependency)):
    """Save text extracted on the Upload page. doc_type is "resume" or "jd"."""
    if payload.doc_type not in VALID_DOC_TYPES:
        raise HTTPException(status_code=400, detail=f"doc_type must be one of {VALID_DOC_TYPES}")
    if not payload.raw_text.strip():
        raise HTTPException(status_code=400, detail="raw_text is empty")
    upload_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": upload_id,
        "user_id": current_user["user_id"],
        "doc_type": payload.doc_type,
        "filename": payload.filename,
        "raw_text": payload.raw_text,
        "char_count": len(payload.raw_text),
        "created_at": now,
    }
    await db.uploaded_texts.insert_one(doc)
    return UploadedTextResponse(**{k: v for k, v in doc.items() if k != "_id"})

@api_router.get("/uploads", response_model=List[UploadedTextResponse])
async def list_uploads(current_user: dict = Depends(get_user_dependency)):
    """List a user's recent uploads (newest first)."""
    docs = await db.uploaded_texts.find(
        {"user_id": current_user["user_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return [UploadedTextResponse(**d) for d in docs]

@api_router.get("/uploads/{upload_id}", response_model=UploadedTextResponse)
async def get_upload(upload_id: str, current_user: dict = Depends(get_user_dependency)):
    doc = await db.uploaded_texts.find_one(
        {"id": upload_id, "user_id": current_user["user_id"]}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Upload not found")
    return UploadedTextResponse(**doc)

# ==================== JD Parse Routes ====================

@api_router.post("/parse/jd/simple")
async def parse_jd_simple(request: ParseRequest):
    """Simple regex-based JD parsing"""
    text = request.text
    
    # Extract job title (usually first line or after common headers)
    job_title = ""
    lines = text.strip().split('\n')
    if lines:
        job_title = lines[0].strip()
    
    # Extract company name
    company_match = re.search(r'(?:company|organization|employer)[:\s]*([^\n]+)', text, re.IGNORECASE)
    company = company_match.group(1).strip() if company_match else ""
    
    # Extract location
    location_match = re.search(r'(?:location|city|place)[:\s]*([^\n]+)', text, re.IGNORECASE)
    location = location_match.group(1).strip() if location_match else ""
    
    # Extract experience required
    exp_match = re.search(r'(\d+[\+\-]?\s*(?:to\s*\d+)?\s*years?)', text, re.IGNORECASE)
    experience_required = exp_match.group(1) if exp_match else ""
    
    # Extract salary range
    salary_match = re.search(r'(?:salary|compensation|pay)[:\s]*([^\n]+)', text, re.IGNORECASE)
    salary_range = salary_match.group(1).strip() if salary_match else ""
    
    # Extract job type
    job_type = "Full-time"  # Default
    if re.search(r'part[\s-]?time', text, re.IGNORECASE):
        job_type = "Part-time"
    elif re.search(r'contract', text, re.IGNORECASE):
        job_type = "Contract"
    elif re.search(r'freelance', text, re.IGNORECASE):
        job_type = "Freelance"
    
    # Extract skills
    skills = []
    skills_section = re.search(r'(?:skills|requirements|qualifications|technologies)[:\s]*([^#]+?)(?=\n\n|\n[A-Z]|$)', text, re.IGNORECASE | re.DOTALL)
    if skills_section:
        skills_text = skills_section.group(1)
        skills = re.split(r'[,\n•\-\|]', skills_text)
        skills = [s.strip() for s in skills if s.strip() and len(s.strip()) < 50]
    
    return {
        "job_title": job_title,
        "company": company,
        "location": location,
        "must_have_skills": skills[:len(skills)//2] if skills else [],
        "good_to_have_skills": skills[len(skills)//2:] if skills else [],
        "experience_required": experience_required,
        "salary_range": salary_range,
        "job_type": job_type,
        "description": text[:500] if len(text) > 500 else text
    }

@api_router.post("/parse/jd/ai")
async def parse_jd_ai(request: ParseRequest):
    """AI-powered JD parsing using Claude"""
    try:
        system_prompt = """You are an expert job description parser and recruiter. Extract structured data from job descriptions.
            Carefully analyze the JD and categorize skills into "Must Have" (required/mandatory) and "Good to Have" (preferred/nice-to-have).

            Return ONLY valid JSON with this exact structure:
            {
                "job_title": "Job Title",
                "company": "Company Name",
                "location": "City, Country or Remote",
                "must_have_skills": ["skill1", "skill2", "skill3"],
                "good_to_have_skills": ["skill1", "skill2"],
                "experience_required": "X-Y years or X+ years",
                "salary_range": "Amount or range if mentioned",
                "job_type": "Full-time/Part-time/Contract/Freelance",
                "description": "Brief job description summary"
            }

            Rules for skill categorization:
            - Must Have: Skills explicitly marked as required, mandatory, essential, must have
            - Good to Have: Skills marked as preferred, nice to have, bonus, plus, advantageous
            - If not explicitly categorized, use context to determine importance

            Extract all information accurately. If something is not found, use empty string or empty array."""
        return await llm_complete_json(
            system_prompt=system_prompt,
            user_prompt=f"Parse this job description and return JSON:\n\n{request.text}",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI JD parsing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"AI parsing failed: {str(e)}")

# ==================== JD CRUD Routes ====================

@api_router.post("/jds", response_model=JDResponse)
async def create_jd(jd_data: JDCreate, current_user: dict = Depends(get_user_dependency)):
    jd_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    jd_doc = {
        "id": jd_id,
        "user_id": current_user["user_id"],
        "title": jd_data.title,
        "raw_text": jd_data.raw_text,
        "parsed_data": jd_data.parsed_data.model_dump() if jd_data.parsed_data else None,
        "created_at": now,
        "updated_at": now
    }
    
    await db.jds.insert_one(jd_doc)
    return JDResponse(**{k: v for k, v in jd_doc.items() if k != "_id"})

@api_router.get("/jds", response_model=List[JDResponse])
async def list_jds(current_user: dict = Depends(get_user_dependency)):
    jds = await db.jds.find(
        {"user_id": current_user["user_id"]}, 
        {"_id": 0}
    ).sort("updated_at", -1).to_list(100)
    return [JDResponse(**j) for j in jds]

@api_router.get("/jds/{jd_id}", response_model=JDResponse)
async def get_jd(jd_id: str, current_user: dict = Depends(get_user_dependency)):
    jd = await db.jds.find_one(
        {"id": jd_id, "user_id": current_user["user_id"]},
        {"_id": 0}
    )
    if not jd:
        raise HTTPException(status_code=404, detail="JD not found")
    return JDResponse(**jd)

@api_router.put("/jds/{jd_id}", response_model=JDResponse)
async def update_jd(jd_id: str, jd_data: JDUpdate, current_user: dict = Depends(get_user_dependency)):
    jd = await db.jds.find_one({"id": jd_id, "user_id": current_user["user_id"]})
    if not jd:
        raise HTTPException(status_code=404, detail="JD not found")
    
    update_dict = {}
    if jd_data.title is not None:
        update_dict["title"] = jd_data.title
    if jd_data.raw_text is not None:
        update_dict["raw_text"] = jd_data.raw_text
    if jd_data.parsed_data is not None:
        update_dict["parsed_data"] = jd_data.parsed_data.model_dump()
    
    update_dict["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.jds.update_one({"id": jd_id}, {"$set": update_dict})
    updated_jd = await db.jds.find_one({"id": jd_id}, {"_id": 0})
    return JDResponse(**updated_jd)

@api_router.delete("/jds/{jd_id}")
async def delete_jd(jd_id: str, current_user: dict = Depends(get_user_dependency)):
    result = await db.jds.delete_one({"id": jd_id, "user_id": current_user["user_id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="JD not found")
    return {"message": "JD deleted successfully"}

# ==================== Telegram Bot Routes ====================

@api_router.get("/telegram/settings")
async def get_telegram_settings(current_user: dict = Depends(get_user_dependency)):
    settings = await db.telegram_settings.find_one(
        {"user_id": current_user["user_id"]},
        {"_id": 0}
    )
    if not settings:
        return {"bot_token": None, "is_active": False, "chat_id": None}
    return settings

@api_router.put("/telegram/settings")
async def update_telegram_settings(settings: TelegramSettingsUpdate, current_user: dict = Depends(get_user_dependency)):
    existing = await db.telegram_settings.find_one({"user_id": current_user["user_id"]})
    
    update_dict = {"user_id": current_user["user_id"]}
    if settings.bot_token is not None:
        update_dict["bot_token"] = settings.bot_token
    if settings.is_active is not None:
        update_dict["is_active"] = settings.is_active
    
    if existing:
        await db.telegram_settings.update_one(
            {"user_id": current_user["user_id"]},
            {"$set": update_dict}
        )
    else:
        update_dict["chat_id"] = None
        await db.telegram_settings.insert_one(update_dict)
    
    # Start or stop bot based on settings
    if settings.is_active and settings.bot_token:
        await start_telegram_bot(settings.bot_token, current_user["user_id"])
    
    updated = await db.telegram_settings.find_one(
        {"user_id": current_user["user_id"]},
        {"_id": 0}
    )
    return updated

# Telegram Bot Handlers
telegram_bots = {}  # Store bot instances by user_id

async def start_telegram_bot(bot_token: str, user_id: str):
    """Start Telegram bot for a user using asyncio"""
    global telegram_bots
    
    try:
        # Stop existing bot for this user if any
        if user_id in telegram_bots:
            old_app = telegram_bots[user_id]
            try:
                await old_app.stop()
                await old_app.shutdown()
            except Exception:
                pass
        
        # Create new application
        application = Application.builder().token(bot_token).build()
        
        # Conversation handler
        conv_handler = ConversationHandler(
            entry_points=[
                CommandHandler("start", lambda u, c: telegram_start(u, c, user_id)),
                MessageHandler(filters.Regex(r'^[Hh]i'), lambda u, c: telegram_start(u, c, user_id))
            ],
            states={
                CHOOSING: [MessageHandler(filters.Regex(r'^(Resume|JD)$'), lambda u, c: telegram_choice(u, c, user_id))],
                WAITING_FOR_TEXT: [MessageHandler(filters.TEXT & ~filters.COMMAND, lambda u, c: telegram_receive_text(u, c, user_id))],
            },
            fallbacks=[CommandHandler("cancel", telegram_cancel)],
        )
        
        application.add_handler(conv_handler)
        
        # Initialize and start
        await application.initialize()
        await application.start()
        
        # Start polling in background task
        asyncio.create_task(run_polling_task(application, user_id))
        
        telegram_bots[user_id] = application
        logger.info(f"Telegram bot started for user {user_id}")
        
    except Exception as e:
        logger.error(f"Failed to start Telegram bot: {str(e)}")
        raise

async def run_polling_task(application, user_id):
    """Run polling as a background task"""
    try:
        updater = application.updater
        await updater.start_polling(drop_pending_updates=True)
        logger.info(f"Polling started for user {user_id}")
    except Exception as e:
        logger.error(f"Polling error for user {user_id}: {str(e)}")

async def telegram_start(update: Update, context: ContextTypes.DEFAULT_TYPE, user_id: str) -> int:
    """Start conversation and ask Resume or JD"""
    reply_keyboard = [["Resume", "JD"]]
    await update.message.reply_text(
        "Hi! I'm your Application Bot. What would you like to parse?\n\n"
        "Choose an option:",
        reply_markup=ReplyKeyboardMarkup(reply_keyboard, one_time_keyboard=True),
    )
    
    # Store chat_id and user_id in context
    context.user_data["app_user_id"] = user_id
    chat_id = str(update.effective_chat.id)
    await db.telegram_settings.update_one(
        {"user_id": user_id},
        {"$set": {"chat_id": chat_id}}
    )
    
    return CHOOSING

async def telegram_choice(update: Update, context: ContextTypes.DEFAULT_TYPE, user_id: str) -> int:
    """Store user's choice and ask for text"""
    choice = update.message.text
    context.user_data["choice"] = choice
    context.user_data["app_user_id"] = user_id
    
    await update.message.reply_text(
        f"Great! Please send me your {choice} text.\n"
        f"I'll parse it and save it for you.",
        reply_markup=ReplyKeyboardRemove(),
    )
    return WAITING_FOR_TEXT

async def telegram_receive_text(update: Update, context: ContextTypes.DEFAULT_TYPE, user_id: str) -> int:
    """Receive and parse the text"""
    text = update.message.text
    choice = context.user_data.get("choice", "Resume")
    
    await update.message.reply_text("Processing your text... Please wait.")
    
    try:
        if choice == "Resume":
            # Parse resume
            parsed = await llm_complete_json(
                system_prompt="""You are a professional resume parser. Extract structured data from resume text.
                Return ONLY valid JSON with: name, email, linkedin, professional_summary, technical_skills, experiences, projects, education.""",
                user_prompt=f"Parse this resume:\n{text}",
            )

            # Save to database
            resume_id = str(uuid.uuid4())
            now = datetime.now(timezone.utc).isoformat()
            await db.resumes.insert_one({
                "id": resume_id,
                "user_id": user_id,
                "title": f"Telegram Resume - {now[:10]}",
                "raw_text": text,
                "parsed_data": parsed,
                "created_at": now,
                "updated_at": now
            })

            await update.message.reply_text(
                f"✅ Resume parsed and saved!\n\n"
                f"Name: {parsed.get('name', 'N/A')}\n"
                f"Skills: {', '.join(parsed.get('technical_skills', [])[:5])}\n\n"
                f"View it in your dashboard!"
            )
        else:
            # Parse JD
            parsed = await llm_complete_json(
                system_prompt="""You are a job description parser. Extract: job_title, company, location, must_have_skills, good_to_have_skills, experience_required, salary_range, job_type.
                Return ONLY valid JSON.""",
                user_prompt=f"Parse this JD:\n{text}",
            )

            # Save to database
            jd_id = str(uuid.uuid4())
            now = datetime.now(timezone.utc).isoformat()
            await db.jds.insert_one({
                "id": jd_id,
                "user_id": user_id,
                "title": f"{parsed.get('job_title', 'JD')} - {now[:10]}",
                "raw_text": text,
                "parsed_data": parsed,
                "created_at": now,
                "updated_at": now
            })

            await update.message.reply_text(
                f"✅ Job Description parsed and saved!\n\n"
                f"Job: {parsed.get('job_title', 'N/A')}\n"
                f"Company: {parsed.get('company', 'N/A')}\n"
                f"Must Have: {', '.join(parsed.get('must_have_skills', [])[:3])}\n\n"
                f"View it in your dashboard!"
            )
    except Exception as e:
        logger.error(f"Telegram parsing error: {str(e)}")
        await update.message.reply_text("Sorry, there was an error parsing your text. Please try again.")
    
    return ConversationHandler.END

async def telegram_cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Cancel the conversation"""
    await update.message.reply_text(
        "Cancelled. Send 'hi' to start again!",
        reply_markup=ReplyKeyboardRemove(),
    )
    return ConversationHandler.END

# ==================== Resume Rewriter Agent ====================

@api_router.post("/rewrite-resume")
async def rewrite_resume(request: RewriteRequest, current_user: dict = Depends(get_user_dependency)):
    """AI-powered resume rewriter using the configured LLM"""
    try:
        # First calculate current score
        try:
            score_result = await llm_complete_json(
                system_prompt="""You are a resume scoring expert. Analyze the resume against the job description.
            Return ONLY a JSON object with a single field "score" which is a number between 0 and 100.
            Score based on: keywords match (40%), skill alignment in context (30%), performance metrics/results (30%).
            Example: {"score": 65}""",
                user_prompt=f"Score this resume against the JD:\n\nRESUME:\n{request.resume_text}\n\nJOB DESCRIPTION:\n{request.jd_text}",
            )
            score_before = score_result.get("score", 50)
        except HTTPException:
            score_before = 50

        # Determine rewrite strategy based on score
        if score_before < 60:
            rewrite_instruction = "FULL REWRITE: The resume score is below 60%. Completely rewrite the resume to better match the JD."
        else:
            rewrite_instruction = "TARGETED IMPROVEMENTS: The resume score is above 60%. Suggest specific improvements for sections that need enhancement."

        # Rewrite resume
        result = await llm_complete_json(
            system_prompt="""You are an expert resume writer and career coach. Your task is to rewrite resumes to match job descriptions.

STRICT RULES FOR REWRITING:
1. KEYWORDS: Always list technical skills and tools as BULLET POINTS, not in sentences
   Example: • Python • React • AWS • Docker

2. KEYWORD ALIGNMENT: When mentioning skills in job descriptions/experience, use them IN CONTEXT within sentences
   Example: "Developed scalable microservices using Python and deployed on AWS infrastructure"

3. RPM (Resume Performance Metrics): ALWAYS include quantifiable results and metrics
   - Revenue/Cost impact: "Reduced costs by $500K annually"
   - Performance gains: "Improved system performance by 40%"
   - Team/Project scale: "Led team of 8 engineers"
   - User impact: "Served 10M+ daily active users"

Return a JSON object with:
{
    "rewritten_resume": "The full rewritten resume text",
    "improvements": ["List of specific improvements made"],
    "estimated_new_score": 85
}""",
            user_prompt=f"{rewrite_instruction}\n\nORIGINAL RESUME:\n{request.resume_text}\n\nTARGET JOB DESCRIPTION:\n{request.jd_text}",
        )

        return {
            "rewritten_resume": result.get("rewritten_resume", ""),
            "improvements": result.get("improvements", []),
            "score_before": score_before,
            "score_after": result.get("estimated_new_score", min(score_before + 20, 95)),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Resume rewrite error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Rewrite failed: {str(e)}")

# ==================== Match Score API ====================

@api_router.post("/match-score", response_model=MatchScoreResponse)
async def calculate_match_score(request: MatchScoreRequest, current_user: dict = Depends(get_user_dependency)):
    """Calculate match score between resume and JD with detailed keyword analysis"""
    try:
        result = await llm_complete_json(
            system_prompt="""You are an expert ATS (Applicant Tracking System) and resume analyzer.

Analyze the resume against the job description and calculate scores based on:

1. KEYWORDS SCORE (0-100):
   - Extract all technical skills, tools, technologies, certifications from JD
   - Check which are present/absent in resume
   - Score = (present keywords / total keywords) * 100

2. KEYWORD ALIGNMENT SCORE (0-100):
   - Check if skills are used IN CONTEXT within experience/project descriptions
   - Not just listed, but demonstrated in sentences with action verbs
   - Higher score for skills shown with impact/results

3. RPM SCORE (Resume Performance Metrics) (0-100):
   - Look for quantifiable achievements: percentages, dollar amounts, time saved
   - Team sizes, user counts, scale of projects
   - Revenue impact, cost savings, efficiency gains

Return ONLY a JSON object:
{
    "keywords_score": 75,
    "alignment_score": 60,
    "rpm_score": 45,
    "present_keywords": ["Python", "React", "AWS"],
    "missing_keywords": ["Kubernetes", "Docker", "CI/CD"],
    "keyword_details": [
        {"keyword": "Python", "present": true, "context": "Used Python for backend development"},
        {"keyword": "Kubernetes", "present": false, "context": null}
    ],
    "suggestions": [
        "Add Kubernetes experience or certification",
        "Include more metrics in your project descriptions"
    ]
}""",
            user_prompt=f"Analyze this resume against the job description:\n\nRESUME:\n{request.resume_text}\n\nJOB DESCRIPTION:\n{request.jd_text}",
        )

        # Calculate overall score (weighted average)
        keywords_score = result.get("keywords_score", 0)
        alignment_score = result.get("alignment_score", 0)
        rpm_score = result.get("rpm_score", 0)
        overall_score = (keywords_score * 0.4) + (alignment_score * 0.3) + (rpm_score * 0.3)

        return MatchScoreResponse(
            overall_score=round(overall_score, 1),
            keywords_score=keywords_score,
            alignment_score=alignment_score,
            rpm_score=rpm_score,
            present_keywords=result.get("present_keywords", []),
            missing_keywords=result.get("missing_keywords", []),
            keyword_details=[KeywordDetail(**kd) for kd in result.get("keyword_details", [])],
            suggestions=result.get("suggestions", [])
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Match score error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Score calculation failed: {str(e)}")

# ==================== Download Routes ====================

@api_router.post("/download/pdf")
async def download_pdf(resume_data: ResumeData):
    """Generate PDF from parsed resume data with professional formatting"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # Narrow margins (0.5 inch = 1.27 cm equivalent)
    left_margin = 0.5 * inch
    right_margin = width - 0.5 * inch
    usable_width = right_margin - left_margin
    
    y = height - 0.5 * inch
    
    def draw_text(text, x, y_pos, font="Helvetica", size=10, color=colors.black):
        c.setFont(font, size)
        c.setFillColor(color)
        c.drawString(x, y_pos, text)
        return y_pos
    
    def draw_right_text(text, y_pos, font="Helvetica", size=10):
        c.setFont(font, size)
        c.drawRightString(right_margin, y_pos, text)
    
    def draw_horizontal_line(y_pos):
        c.setStrokeColor(colors.black)
        c.setLineWidth(0.5)
        c.line(left_margin, y_pos, right_margin, y_pos)
        return y_pos - 0.1 * inch
    
    def wrap_text(text, max_width, font, size):
        c.setFont(font, size)
        words = text.split()
        lines = []
        current_line = ""
        for word in words:
            test_line = f"{current_line} {word}".strip()
            if c.stringWidth(test_line, font, size) < max_width:
                current_line = test_line
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word
        if current_line:
            lines.append(current_line)
        return lines
    
    def check_page_break(y_pos, needed=1.5):
        if y_pos < needed * inch:
            c.showPage()
            return height - 0.5 * inch
        return y_pos
    
    # Name
    if resume_data.name:
        y = draw_text(resume_data.name, left_margin, y, "Helvetica-Bold", 16)
        y -= 0.22 * inch
    
    # Contact Info
    contact_parts = []
    if resume_data.email:
        contact_parts.append(resume_data.email)
    if resume_data.linkedin:
        contact_parts.append(resume_data.linkedin)
    if contact_parts:
        y = draw_text(" | ".join(contact_parts), left_margin, y, "Helvetica", 10, colors.HexColor("#475569"))
        y -= 0.35 * inch
    
    # Professional Summary
    if resume_data.professional_summary:
        y = draw_text("PROFESSIONAL SUMMARY", left_margin, y, "Helvetica-Bold", 14)
        y -= 0.05 * inch
        y = draw_horizontal_line(y)
        y -= 0.1 * inch
        lines = wrap_text(resume_data.professional_summary, usable_width, "Helvetica", 12)
        for line in lines:
            y = draw_text(line, left_margin, y, "Helvetica", 12)
            y -= 0.2 * inch
        y -= 0.15 * inch
    
    # Technical Skills
    if resume_data.technical_skills:
        y = check_page_break(y)
        y = draw_text("TECHNICAL SKILLS", left_margin, y, "Helvetica-Bold", 14)
        y -= 0.05 * inch
        y = draw_horizontal_line(y)
        y -= 0.1 * inch
        skills_text = "Skills: " + ", ".join(resume_data.technical_skills)
        lines = wrap_text(skills_text, usable_width, "Helvetica", 12)
        for line in lines:
            y = draw_text(line, left_margin, y, "Helvetica", 12)
            y -= 0.2 * inch
        y -= 0.15 * inch
    
    # Experience
    if resume_data.experiences:
        y = check_page_break(y)
        y = draw_text("PROFESSIONAL EXPERIENCE", left_margin, y, "Helvetica-Bold", 14)
        y -= 0.05 * inch
        y = draw_horizontal_line(y)
        y -= 0.1 * inch
        
        for exp in resume_data.experiences:
            y = check_page_break(y)
            title = exp.get("title", "")
            company = exp.get("company", "")
            from_date = exp.get("from_date", "")
            to_date = exp.get("to_date", "")
            date_str = f"{from_date} - {to_date}".strip(' -') if from_date or to_date else ""
            
            # Title - Company on left, Date on right
            left_text = f"{title} - {company}" if company else title
            y = draw_text(left_text, left_margin, y, "Helvetica", 14)
            if date_str:
                draw_right_text(date_str, y, "Helvetica", 12)
            y -= 0.22 * inch
            
            for bullet in exp.get("bullets", []):
                y = check_page_break(y)
                lines = wrap_text(f"• {bullet}", usable_width - 0.25 * inch, "Helvetica", 12)
                for line in lines:
                    y = draw_text(line, left_margin + 0.15 * inch, y, "Helvetica", 12)
                    y -= 0.18 * inch
            y -= 0.1 * inch
        y -= 0.1 * inch
    
    # Projects
    if resume_data.projects:
        y = check_page_break(y)
        y = draw_text("PROJECTS", left_margin, y, "Helvetica-Bold", 14)
        y -= 0.05 * inch
        y = draw_horizontal_line(y)
        y -= 0.1 * inch
        
        for proj in resume_data.projects:
            y = check_page_break(y)
            title = proj.get("title", "")
            org = proj.get("organization", "") or proj.get("description", "")
            from_date = proj.get("from_date", "")
            to_date = proj.get("to_date", "")
            date_str = f"{from_date} - {to_date}".strip(' -') if from_date or to_date else ""
            
            left_text = f"{title} - {org}" if org else title
            y = draw_text(left_text, left_margin, y, "Helvetica", 14)
            if date_str:
                draw_right_text(date_str, y, "Helvetica", 12)
            y -= 0.22 * inch
            
            for bullet in proj.get("bullets", []):
                y = check_page_break(y)
                lines = wrap_text(f"• {bullet}", usable_width - 0.25 * inch, "Helvetica", 12)
                for line in lines:
                    y = draw_text(line, left_margin + 0.15 * inch, y, "Helvetica", 12)
                    y -= 0.18 * inch
            y -= 0.1 * inch
        y -= 0.1 * inch
    
    # Education
    if resume_data.education:
        y = check_page_break(y)
        y = draw_text("EDUCATION", left_margin, y, "Helvetica-Bold", 14)
        y -= 0.05 * inch
        y = draw_horizontal_line(y)
        y -= 0.1 * inch
        
        for edu in resume_data.education:
            y = check_page_break(y)
            title = edu.get("title", "")
            institution = edu.get("institution", "")
            from_date = edu.get("from_date", "")
            to_date = edu.get("to_date", "")
            date_str = f"{from_date} - {to_date}".strip(' -') if from_date or to_date else ""
            
            if institution:
                y = draw_text(institution, left_margin, y, "Helvetica-Bold", 12)
                y -= 0.2 * inch
            
            line = title
            if date_str:
                line = f"{title} | {date_str}"
            y = draw_text(line, left_margin, y, "Helvetica", 12)
            y -= 0.2 * inch
    
    c.save()
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=resume.pdf"}
    )

@api_router.post("/download/docx")
async def download_docx(resume_data: ResumeData):
    """Generate DOCX from parsed resume data with professional formatting"""
    doc = Document()
    
    # Helper function to add horizontal line under section headers
    def add_horizontal_line(paragraph):
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    # Helper function to calculate tab stop for right-aligned dates
    def calculate_tab_stop(section, margin_from_edge_cm=0.2, page_width_cm=21.59):
        left_cm = section.left_margin.cm
        right_cm = section.right_margin.cm
        usable_cm = page_width_cm - left_cm - right_cm
        tab_cm = usable_cm - margin_from_edge_cm
        tab_in = tab_cm / 2.54
        return Inches(tab_in)
    
    # Helper function to add experience/project entry with right-aligned date
    def add_entry_with_date(doc, section, title, subtitle, date_str, title_size=14, subtitle_size=14):
        para = doc.add_paragraph()
        # Title
        title_run = para.add_run(title)
        title_run.font.size = Pt(title_size)
        title_run.font.bold = False
        
        if subtitle:
            sep = para.add_run(' - ')
            sep.font.size = Pt(title_size)
            sub_run = para.add_run(subtitle)
            sub_run.font.size = Pt(subtitle_size)
            sub_run.font.bold = False
        
        if date_str:
            # Calculate and add tab stop for right alignment
            tab_pos = calculate_tab_stop(section, margin_from_edge_cm=0.2)
            para.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
            date_run = para.add_run('\t' + date_str)
            date_run.font.size = Pt(12)
        
        para.paragraph_format.left_indent = Inches(0.0)
        return para
    
    # Set default document font to Aptos
    try:
        style = doc.styles['Normal']
        style.font.name = 'Aptos'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
    except Exception:
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
    
    # Set narrow margins (1.27 cm)
    sections = doc.sections
    section = sections[0]
    for sec in sections:
        sec.top_margin = Cm(1.27)
        sec.bottom_margin = Cm(1.27)
        sec.left_margin = Cm(1.27)
        sec.right_margin = Cm(1.27)
    
    # Name Header
    if resume_data.name:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(resume_data.name)
        name_run.bold = True
        name_run.font.size = Pt(16)
    
    # Contact Info
    contact_parts = []
    if resume_data.email:
        contact_parts.append(resume_data.email)
    if resume_data.linkedin:
        contact_parts.append(resume_data.linkedin)
    if contact_parts:
        contact_para = doc.add_paragraph(' | '.join(contact_parts))
        if contact_para.runs:
            contact_para.runs[0].font.size = Pt(10)
    
    # Professional Summary
    if resume_data.professional_summary:
        sh = doc.add_paragraph()
        sh_run = sh.add_run('PROFESSIONAL SUMMARY')
        sh_run.font.size = Pt(14)
        sh_run.bold = True
        add_horizontal_line(sh)
        doc.add_paragraph(resume_data.professional_summary)
    
    # Technical Skills
    if resume_data.technical_skills:
        sk = doc.add_paragraph()
        sk_run = sk.add_run('TECHNICAL SKILLS')
        sk_run.font.size = Pt(14)
        sk_run.bold = True
        add_horizontal_line(sk)
        
        skills_para = doc.add_paragraph()
        skills_run = skills_para.add_run('Skills: ')
        skills_run.bold = True
        skills_para.add_run(', '.join(resume_data.technical_skills))
        
        doc.add_paragraph()
    
    # Professional Experience
    if resume_data.experiences:
        eh = doc.add_paragraph()
        eh_run = eh.add_run('PROFESSIONAL EXPERIENCE')
        eh_run.font.size = Pt(14)
        eh_run.bold = True
        add_horizontal_line(eh)
        
        for exp in resume_data.experiences:
            title = exp.get("title", "")
            company = exp.get("company", "")
            from_date = exp.get("from_date", "")
            to_date = exp.get("to_date", "")
            date_str = f"{from_date} - {to_date}".strip(' -') if from_date or to_date else ""
            
            add_entry_with_date(doc, section, title, company, date_str)
            
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style='List Bullet')
        
        doc.add_paragraph()
    
    # Projects
    if resume_data.projects:
        pj = doc.add_paragraph()
        pj_run = pj.add_run('PROJECTS')
        pj_run.font.size = Pt(14)
        pj_run.bold = True
        add_horizontal_line(pj)
        
        for proj in resume_data.projects:
            title = proj.get("title", "")
            organization = proj.get("organization", "") or proj.get("description", "")
            from_date = proj.get("from_date", "")
            to_date = proj.get("to_date", "")
            date_str = f"{from_date} - {to_date}".strip(' -') if from_date or to_date else ""
            
            add_entry_with_date(doc, section, title, organization, date_str)
            
            for bullet in proj.get("bullets", []):
                doc.add_paragraph(bullet, style='List Bullet')
        
        doc.add_paragraph()
    
    # Education
    if resume_data.education:
        ed = doc.add_paragraph()
        ed_run = ed.add_run('EDUCATION')
        ed_run.font.size = Pt(14)
        ed_run.bold = True
        add_horizontal_line(ed)
        
        for edu in resume_data.education:
            title = edu.get("title", "")
            institution = edu.get("institution", "")
            from_date = edu.get("from_date", "")
            to_date = edu.get("to_date", "")
            
            if institution:
                inst_para = doc.add_paragraph()
                inst_run = inst_para.add_run(institution)
                inst_run.bold = True
            
            date_str = f"{from_date} - {to_date}".strip(' -') if from_date or to_date else ""
            line = title
            if date_str:
                line = f"{title} | {date_str}"
            doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=resume.docx"}
    )

# ==================== Health Check ====================

@api_router.get("/")
async def root():
    return {"message": "Resume Formatter API", "status": "healthy"}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def init_auth_indexes():
    """Create auth-related indexes and sweep expired sessions at startup.

    * Unique index on db.users.email prevents two simultaneous register
      calls from inserting duplicate accounts.
    * TTL index on db.user_sessions.expires_at makes MongoDB itself
      garbage-collect expired sessions, so we never accumulate dead rows.
    """
    try:
        await db.users.create_index("email", unique=True)
    except Exception as e:  # noqa: BLE001
        logger.warning(f"Could not ensure unique email index: {e}")
    try:
        # TTL index — Mongo deletes the doc when expires_at passes.
        await db.user_sessions.create_index("expires_at", expireAfterSeconds=0)
    except Exception as e:  # noqa: BLE001
        logger.warning(f"Could not ensure sessions TTL index: {e}")


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
